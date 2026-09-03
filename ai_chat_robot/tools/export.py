"""导出工具：CSV / XLSX / DOCX（统一版式）。"""

from __future__ import annotations

import csv
import json
import re
from pathlib import Path

from agents import function_tool

from config.file_agent import FILE_AGENT_MAX_WRITE_BYTES, FILE_AGENT_WORKSPACE
from guardrails import validate_file_tool_path
from sandbox.audit import log_audit_event
from tools.file import CSV_ENCODING, resolve_safe_path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:  # pragma: no cover
    Document = None  # type: ignore[misc, assignment]
    WD_ALIGN_PARAGRAPH = None  # type: ignore[misc, assignment]
    qn = None  # type: ignore[misc, assignment]
    Cm = None  # type: ignore[misc, assignment]
    Pt = None  # type: ignore[misc, assignment]

try:
    from openpyxl import Workbook
except ImportError:  # pragma: no cover
    Workbook = None  # type: ignore[misc, assignment]

_BODY_FONT = "宋体"
_TITLE_FONT = "黑体"
_BODY_SIZE_PT = 12
_TITLE_SIZE_PT = 16


def _safe_csv_cell(value: object) -> object:
    """Prevent spreadsheet formula execution when opening exported CSVs."""
    text = str(value)
    if text.startswith(("=", "+", "-", "@")):
        return "'" + text
    return value


def _parse_table_json(table_json: str) -> tuple[list[str], list[list]]:
    data = json.loads(table_json)
    if isinstance(data, dict) and "headers" in data and "rows" in data:
        headers = [str(h) for h in data["headers"]]
        rows = [[str(cell) for cell in row] for row in data["rows"]]
        return headers, rows
    if isinstance(data, list) and data and isinstance(data[0], dict):
        headers = list(data[0].keys())
        rows = [[str(item.get(h, "")) for h in headers] for item in data]
        return headers, rows
    raise ValueError("table_json 需为 {headers, rows} 或对象数组")


def _strip_markdown_line(line: str) -> str:
    text = line.strip()
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^>\s*", "", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text.strip()


def _set_run_font(run, *, font_name: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if qn is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _format_docx_path_message(relative_path: str, path: Path, size: int) -> str:
    return (
        f"已导出 DOCX。\n"
        f"- 相对路径：`{relative_path}`\n"
        f"- 本机绝对路径：`{path}`\n"
        f"- 大小：{size} bytes\n"
        f"请在资源管理器打开：`{FILE_AGENT_WORKSPACE}` 下的 `{relative_path}`"
    )


def export_table_csv_impl(relative_path: str, table_json: str) -> str:
    headers, rows = _parse_table_json(table_json)
    path = resolve_safe_path(relative_path, create_parents=True)
    with path.open("w", encoding=CSV_ENCODING, newline="") as handle:
        writer = csv.writer(handle)
        writer.writerow([_safe_csv_cell(cell) for cell in headers])
        writer.writerows(
            [[_safe_csv_cell(cell) for cell in row] for row in rows]
        )
    size = path.stat().st_size
    log_audit_event(
        "file_agent_export_csv",
        status="ok",
        detail=relative_path,
        extra={"rows": len(rows), "bytes": size},
    )
    return f"已导出 CSV：`{relative_path}`（{len(rows)} 行，{size} bytes）。"


def export_table_xlsx_impl(relative_path: str, table_json: str, sheet_name: str = "Sheet1") -> str:
    if Workbook is None:
        return "缺少 openpyxl，请执行：pip install openpyxl"
    headers, rows = _parse_table_json(table_json)
    path = resolve_safe_path(relative_path, create_parents=True)
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "Sheet1"
    ws.append(headers)
    for row in rows:
        ws.append(row)
    wb.save(path)
    size = path.stat().st_size
    log_audit_event(
        "file_agent_export_xlsx",
        status="ok",
        detail=relative_path,
        extra={"rows": len(rows), "bytes": size},
    )
    return f"已导出 XLSX：`{relative_path}`（{len(rows)} 行，{size} bytes）。"


def export_docx_impl(relative_path: str, title: str, body: str) -> str:
    if Document is None:
        return "缺少 python-docx，请执行：pip install python-docx"

    if len(body.encode("utf-8")) > FILE_AGENT_MAX_WRITE_BYTES:
        return (
            f"正文过长（>{FILE_AGENT_MAX_WRITE_BYTES} bytes），请缩短或分段导出。"
        )

    path = resolve_safe_path(relative_path, create_parents=True)
    doc = Document()

    # 正文默认样式：宋体 12 磅
    normal = doc.styles["Normal"]
    normal.font.name = _BODY_FONT
    normal.font.size = Pt(_BODY_SIZE_PT)
    if qn is not None:
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), _BODY_FONT)

    clean_title = _strip_markdown_line(title)
    if clean_title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(clean_title)
        _set_run_font(title_run, font_name=_TITLE_FONT, size_pt=_TITLE_SIZE_PT, bold=True)
        title_para.paragraph_format.space_after = Pt(12)

    normalized = body.replace("\r\n", "\n").strip()
    blocks = re.split(r"\n\s*\n", normalized) if "\n\n" in normalized else normalized.split("\n")

    for block in blocks:
        for line in block.split("\n"):
            text = _strip_markdown_line(line)
            if not text:
                continue
            para = doc.add_paragraph()
            run = para.add_run(text)
            _set_run_font(run, font_name=_BODY_FONT, size_pt=_BODY_SIZE_PT)
            para.paragraph_format.line_spacing = 1.5
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.first_line_indent = Cm(0.74)

    doc.save(path)
    size = path.stat().st_size
    log_audit_event(
        "file_agent_export_docx",
        status="ok",
        detail=relative_path,
        extra={"bytes": size},
    )
    return _format_docx_path_message(relative_path, path, size)


@function_tool(tool_input_guardrails=[validate_file_tool_path])
def export_table_csv(relative_path: str, table_json: str) -> str:
    """将表格 JSON 导出为 CSV（UTF-8 BOM）。table_json: {headers:[], rows:[][]} 或对象数组。"""
    return export_table_csv_impl(relative_path, table_json)


@function_tool(tool_input_guardrails=[validate_file_tool_path])
def export_table_xlsx(
    relative_path: str,
    table_json: str,
    sheet_name: str = "Sheet1",
) -> str:
    """将表格 JSON 导出为 Excel xlsx。"""
    return export_table_xlsx_impl(relative_path, table_json, sheet_name=sheet_name)


@function_tool(tool_input_guardrails=[validate_file_tool_path])
def export_docx(relative_path: str, title: str, body: str) -> str:
    """将标题与正文导出为 Word docx。body 须为纯文本（勿用 Markdown 符号），按行或空行分段。"""
    return export_docx_impl(relative_path, title, body)
